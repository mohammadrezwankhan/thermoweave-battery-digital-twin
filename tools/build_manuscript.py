"""Build the journal-style ThermoWeave-3D Word manuscript.

Design preset: narrative_proposal with editorial_cover from the Codex
documents skill. Token map: US Letter; 1-inch margins; Calibri 11 pt body;
1.333 line spacing; 8 pt paragraph spacing; dark navy headings; restrained
blue accent; editorial running header; bottom page number and draft status.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "manuscript.md"
OUTPUT = ROOT / "manuscript" / "ThermoWeave_3D_Manuscript.docx"

NAVY = "102A43"
BLUE = "1769AA"
PALE_BLUE = "EAF3F8"
MID_GREY = "5C6873"
LIGHT_GREY = "D9E2E8"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_repeat_table_header(row) -> None:
    repeat_table_header(row)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, value, end):
        run._r.append(element)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(https?://[^\s]+|`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            url = token.rstrip(".,")
            suffix = token[len(url) :]
            add_hyperlink(paragraph, url, url)
            if suffix:
                paragraph.add_run(suffix)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1D2730")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    for list_style_name in ("List Number", "List Bullet"):
        list_style = styles[list_style_name]
        list_style.font.name = "Calibri"
        list_style.font.size = Pt(11)
        list_style.paragraph_format.space_after = Pt(3)
        list_style.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Title": (20, 0, 16, NAVY),
        "Heading 1": (16, 18, 10, NAVY),
        "Heading 2": (13, 14, 6, NAVY),
        "Heading 3": (11, 10, 4, BLUE),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, size, color in (
        ("Caption", 9, MID_GREY),
        ("Subtitle", 11, MID_GREY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = "Cambria Math"
    equation.font.size = Pt(10.5)
    equation.font.color.rgb = RGBColor.from_string(NAVY)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(7)
    equation.paragraph_format.space_after = Pt(7)


def configure_section(section, first=False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = first


def add_running_matter(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = "THERMOWEAVE-3D  •  METHODS MANUSCRIPT"
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_repeat_table_header(table.rows[0])
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.3)
    left = table.cell(0, 0).paragraphs[0]
    left.add_run("JOURNAL-FORMATTED DRAFT • NOT PEER REVIEWED")
    left.style = "Caption"
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run("Page ")
    add_page_field(right)


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("REPRODUCIBLE BATTERY THERMAL RESEARCH SOFTWARE")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "ThermoWeave-3D: A deterministic Cartesian graph electrothermal model "
        "for layer-resolved battery-module studies"
    )
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(subtitle, "Mohammad Rezwan Khan")
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(13)
    for line in (
        "Methods and research-software article",
        "Version 0.2.0-draft • 14 August 2026",
        "Affiliation and correspondence to be confirmed before submission",
    ):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = "Subtitle"

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(box.rows[0])
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(5.6)
    cell = box.cell(0, 0)
    shade(cell, PALE_BLUE)
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Submission boundary: journal-formatted preprint draft; not peer reviewed; "
        "synthetic evidence only; human scientific review required before submission."
    )
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_paragraph()
    visual = ROOT / "artifacts" / "figures" / "3d-module-temperature.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(visual), width=Inches(4.25))
    set_alt_text(
        shape,
        "ThermoWeave-3D cover visual",
        "Three-layer by twelve-position synthetic battery graph colored by final node temperature.",
    )
    p = doc.add_paragraph("Independent implementation • exact scenarios • traceable artifacts")
    p.style = "Caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_section(WD_SECTION.NEW_PAGE)


def is_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", line))


def parse_table(lines, start):
    table_lines = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = [[part.strip() for part in row.strip("|").split("|")] for row in table_lines]
    if len(rows) > 1 and is_separator(table_lines[1]):
        rows.pop(1)
    return rows, index


def add_table(doc: Document, rows) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_index, values in enumerate(rows):
        row = table.rows[r_index]
        if r_index == 0:
            set_repeat_table_header(row)
        for c_index, value in enumerate(values):
            cell = row.cells[c_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_index == 0:
                shade(cell, NAVY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if r_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            if r_index % 2 == 0 and r_index != 0:
                shade(cell, "F4F7F9")
    doc.add_paragraph()


def add_markdown_image(doc: Document, line: str) -> None:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    alt, path = match.groups()
    image_path = (SOURCE.parent / path).resolve()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(6.35 if "layer-response" in path else 5.8)
    shape = paragraph.add_run().add_picture(str(image_path), width=width)
    set_alt_text(shape, alt[:80] or "Manuscript figure", alt or "Scientific figure")


def build() -> Path:
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0], first=True)
    document.core_properties.title = "ThermoWeave-3D"
    document.core_properties.subject = "Deterministic 3-D battery electrothermal graph model"
    document.core_properties.author = "Mohammad Rezwan Khan"
    document.core_properties.keywords = (
        "battery thermal model, graph Laplacian, electrothermal simulation, reproducibility"
    )
    document.core_properties.comments = (
        "Journal-formatted draft generated with AI assistance; human review required."
    )
    add_cover(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines = []
    skipped_title = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not line:
            index += 1
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if line.startswith("**Mohammad") or line.startswith("Affiliation and") or line.startswith("**Manuscript type") or line.startswith("**Version") or line.startswith("**Status"):
            index += 1
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("# "):
            document.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith("!["):
            add_markdown_image(document, line)
        elif line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        elif line.startswith("> "):
            p = document.add_paragraph(style="Equation")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\.\s+", line):
            p = document.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.startswith("**Figure") or line.startswith("**Table"):
            p = document.add_paragraph(style="Caption")
            add_inline(p, line)
        elif line == "## Abstract":
            document.add_paragraph("Abstract", style="Heading 1")
        else:
            p = document.add_paragraph()
            add_inline(p, line)
        index += 1

    for section_index, section in enumerate(document.sections):
        configure_section(section, first=section_index == 0)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        if section_index == 0:
            section.header.paragraphs[0].text = ""
            section.footer.paragraphs[0].text = ""
        else:
            add_running_matter(section)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
